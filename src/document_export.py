"""Read-only, in-memory Word and PDF export for saved BRDs and PRDs."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from enum import Enum
from html import escape
from io import BytesIO
from pathlib import Path
import re
import unicodedata
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)

from src.agile import AgileArtifactType
from src.document_templates import document_template
from src.models import DocumentStatus, DocumentType, Product, ProductDocument


class DocumentExportError(ValueError):
    """A user-safe export error without path, provider, or secret detail."""


class ExportFormat(str, Enum):
    DOCX = "docx"
    PDF = "pdf"


@dataclass(frozen=True)
class ExportResult:
    content: bytes
    filename: str
    mime_type: str
    generated_at: datetime


@dataclass(frozen=True)
class ExportBlock:
    kind: str
    title: str
    content: str | tuple[tuple[str, ...], ...]
    widths: tuple[float, ...] = ()


INK = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_FILL = "F2F4F7"
def _utc_time(value: datetime | None) -> datetime:
    timestamp = value or datetime.now(timezone.utc)
    if timestamp.tzinfo is None:
        raise DocumentExportError("Export time must include a timezone.")
    return timestamp.astimezone(timezone.utc)


def _display_time(value: datetime) -> str:
    return value.strftime("%Y-%m-%d %H:%M UTC")


def _safe_component(value: object, *, fallback: str, maximum: int = 48) -> str:
    normalized = unicodedata.normalize("NFKD", str(value or ""))
    ascii_value = normalized.encode("ascii", "ignore").decode("ascii").lower()
    component = re.sub(r"[^a-z0-9]+", "-", ascii_value).strip("-.")
    component = component[:maximum].rstrip("-.")
    return component or fallback


def export_filename(
    product: Product,
    document: ProductDocument,
    export_format: ExportFormat | str,
) -> str:
    """Return a deterministic, descriptive, path-safe export filename."""

    try:
        selected_format = ExportFormat(export_format)
    except (TypeError, ValueError) as error:
        raise DocumentExportError("Choose Word or PDF export.") from error
    if product.id is None or document.id is None or document.product_id != product.id:
        raise DocumentExportError("Only a saved document associated with this product can be exported.")
    product_name = _safe_component(product.name, fallback=f"product-{product.id}")
    version = _safe_component(document.version, fallback="version", maximum=24)
    filename = (
        f"pmc-{product_name}-{document.document_type.value.lower()}-"
        f"{document.id}-v{version}.{selected_format.value}"
    )
    return filename[:180]


def _clean(value: object, *, blank: str = "Not provided") -> str:
    text = str(value or "").strip()
    return text or blank


def _status_label(document: ProductDocument) -> str:
    return "Approved" if document.document_status is DocumentStatus.APPROVED else "Draft"


def _legacy_is_represented(section_text: str, values: tuple[str, ...], legacy_ids: tuple[str, ...]) -> bool:
    if any(identifier.startswith("legacy-") for identifier in legacy_ids):
        return True
    normalized = " ".join(section_text.split()).casefold()
    return bool(normalized) and normalized in {
        " ".join(value.split()).casefold() for value in values if value.strip()
    }


def _legacy_block(
    *,
    label: str,
    section_text: str,
    values: tuple[str, ...],
    legacy_ids: tuple[str, ...],
) -> ExportBlock | None:
    if not section_text.strip() or _legacy_is_represented(section_text, values, legacy_ids):
        return None
    return ExportBlock("text", f"{label} — preserved legacy content", section_text)


def _criteria_text(criteria: object) -> str:
    entries = tuple(criteria)
    if not entries:
        return "Not provided"
    return "\n".join(f"{item.position}. {_clean(item.text)}" for item in entries)


def _prd_contributor_blocks(document: ProductDocument) -> list[ExportBlock]:
    rows = tuple(
        (_clean(row.contributor_name), _clean(row.contributor_role))
        for row in document.contributors
    )
    blocks = [
        ExportBlock(
            "table",
            "Contributors and Roles",
            (("Contributor", "Role"), *rows) if rows else (("Contributor", "Role"), ("Not provided", "Not provided")),
            (2.7, 3.8),
        )
    ]
    legacy = _legacy_block(
        label="Contributors and roles",
        section_text=document.sections.get("contributors_roles", ""),
        values=tuple(value for row in rows for value in row),
        legacy_ids=tuple(row.entry_id for row in document.contributors),
    )
    if legacy:
        blocks.append(legacy)
    return blocks


def _prd_milestone_blocks(document: ProductDocument) -> list[ExportBlock]:
    rows = tuple((_clean(row.date), _clean(row.milestone)) for row in document.key_dates_milestones)
    blocks = [
        ExportBlock(
            "table",
            "Key Dates and Milestones",
            (("Date", "Milestone"), *rows) if rows else (("Date", "Milestone"), ("Not provided", "Not provided")),
            (1.6, 4.9),
        )
    ]
    legacy_date = _legacy_block(
        label="Key dates",
        section_text=document.sections.get("key_dates", ""),
        values=tuple(row[0] for row in rows),
        legacy_ids=tuple(row.entry_id for row in document.key_dates_milestones),
    )
    legacy_milestone = _legacy_block(
        label="Milestones",
        section_text=document.sections.get("milestones", ""),
        values=tuple(row[1] for row in rows),
        legacy_ids=tuple(row.entry_id for row in document.key_dates_milestones),
    )
    blocks.extend(block for block in (legacy_date, legacy_milestone) if block)
    return blocks


def _prd_hierarchy_blocks(document: ProductDocument) -> list[ExportBlock]:
    blocks = [ExportBlock("heading", "Structured Agile Hierarchy", "")]
    if not document.agile_hierarchy:
        blocks.append(ExportBlock("text", "Hierarchy", "No structured hierarchy entries provided."))
        return blocks
    counts = {artifact_type: 0 for artifact_type in AgileArtifactType}
    for artifact in document.agile_hierarchy:
        counts[artifact.artifact_type] += 1
        parent = artifact.parent_artifact_id or "None"
        blocks.append(
            ExportBlock(
                "table",
                f"{artifact.artifact_type.value.replace('_', ' ').title()} {artifact.position}: {_clean(artifact.title, blank='Untitled')}",
                (
                    ("Field", "Saved content"),
                    ("Stable ID", artifact.artifact_id),
                    ("Parent ID", parent),
                    ("Description", _clean(artifact.description)),
                    ("Acceptance criteria", _criteria_text(artifact.acceptance_criteria)),
                ),
                (1.6, 4.9),
            )
        )
    blocks.append(
        ExportBlock(
            "text",
            "Informational hierarchy summary",
            " · ".join(
                (
                    f"Epics: {counts[AgileArtifactType.EPIC]}",
                    f"Capabilities: {counts[AgileArtifactType.CAPABILITY]}",
                    f"Features: {counts[AgileArtifactType.FEATURE]}",
                    f"User Stories: {counts[AgileArtifactType.USER_STORY]}",
                    f"Acceptance criteria: {sum(len(row.acceptance_criteria) for row in document.agile_hierarchy)}",
                )
            ),
        )
    )
    return blocks


def _success_matrix_blocks(document: ProductDocument) -> list[ExportBlock]:
    blocks = [ExportBlock("heading", "PRD Success Matrix", "")]
    if not document.success_matrix:
        blocks.append(ExportBlock("text", "Success Matrix", "No Success Matrix entries provided."))
        return blocks
    for entry in document.success_matrix:
        status = entry.status.value.replace("_", " ").title() if entry.status else "Not provided"
        blocks.append(
            ExportBlock(
                "table",
                f"Outcome {entry.position}: {_clean(entry.requirement_outcome)}",
                (
                    ("Field", "Saved content"),
                    ("Stable ID", entry.entry_id),
                    ("Metric", _clean(entry.metric)),
                    ("Baseline", _clean(entry.baseline, blank="Not known")),
                    ("Target", _clean(entry.target)),
                    ("Minimum acceptance threshold", _clean(entry.minimum_acceptance_threshold)),
                    ("Measurement method", _clean(entry.measurement_method)),
                    ("Data source", _clean(entry.data_source)),
                    ("Evaluation period", _clean(entry.evaluation_period)),
                    ("Validation owner", _clean(entry.validation_owner)),
                    ("Status", status),
                ),
                (2.1, 4.4),
            )
        )
    return blocks


def _brd_hierarchy_blocks(document: ProductDocument) -> list[ExportBlock]:
    blocks = [ExportBlock("heading", "BRD Agile Hierarchy", "")]
    if not document.brd_hierarchy:
        blocks.append(ExportBlock("text", "Hierarchy", "No structured BRD hierarchy rows provided."))
        return blocks
    for row in document.brd_hierarchy:
        data: list[tuple[str, str]] = [("Field", "Saved content")]
        ids: list[str] = [row.row_id]
        values: list[str] = []
        for level in ("epic", "capability", "feature", "user_story"):
            label = level.replace("_", " ").title()
            value = _clean(getattr(row, level))
            values.append(value)
            ids.append(getattr(row, f"{level}_id"))
            data.append((label, value))
            data.append((f"{label} acceptance criteria", _criteria_text(getattr(row, f"{level}_acceptance_criteria"))))
        blocks.append(ExportBlock("table", f"Hierarchy row {row.position}", tuple(data), (2.1, 4.4)))
        for key, label, value in (
            ("epics", "Epics", values[0]),
            ("capabilities", "Capabilities", values[1]),
            ("features", "Features", values[2]),
            ("user_stories", "User stories", values[3]),
            ("acceptance_criteria", "Associated acceptance criteria", _criteria_text(row.user_story_acceptance_criteria)),
        ):
            legacy = _legacy_block(
                label=label,
                section_text=document.sections.get(key, ""),
                values=(value,),
                legacy_ids=tuple(ids),
            )
            if legacy:
                blocks.append(legacy)
    return blocks


def _brd_risk_blocks(document: ProductDocument) -> list[ExportBlock]:
    rows = tuple((_clean(row.business_risk), _clean(row.mitigation_strategy)) for row in document.brd_risks)
    blocks = [
        ExportBlock(
            "table",
            "Business Risk and Mitigation Strategy",
            (("Business risk", "Mitigation strategy"), *rows)
            if rows else (("Business risk", "Mitigation strategy"), ("Not provided", "Not provided")),
            (3.0, 3.5),
        )
    ]
    for key, label, index in (
        ("business_risks", "Business risks", 0),
        ("mitigation_strategies", "Mitigation strategies", 1),
    ):
        legacy = _legacy_block(
            label=label,
            section_text=document.sections.get(key, ""),
            values=tuple(row[index] for row in rows),
            legacy_ids=tuple(row.entry_id for row in document.brd_risks),
        )
        if legacy:
            blocks.append(legacy)
    return blocks


def export_blocks(document: ProductDocument) -> tuple[ExportBlock, ...]:
    """Return one authoritative ordered content model shared by DOCX and PDF."""

    hidden = {
        DocumentType.PRD: {"contributors_roles", "key_dates", "milestones"},
        DocumentType.BRD: {
            "epics", "capabilities", "features", "user_stories",
            "acceptance_criteria", "business_risks", "mitigation_strategies",
        },
    }
    blocks: list[ExportBlock] = []
    definitions = document_template(document.document_type)
    current_group: str | None = None
    for index, definition in enumerate(definitions):
        if definition.group != current_group:
            if current_group == "4. Functional Requirements" and document.document_type is DocumentType.PRD:
                blocks.extend(_prd_hierarchy_blocks(document))
            if current_group == "5. Business Requirements" and document.document_type is DocumentType.BRD:
                blocks.extend(_brd_hierarchy_blocks(document))
            if current_group == "7. Success Metrics and KPIs" and document.document_type is DocumentType.PRD:
                blocks.extend(_success_matrix_blocks(document))
            current_group = definition.group
            blocks.append(ExportBlock("group", current_group, ""))
        if document.document_type is DocumentType.PRD and definition.key == "contributors_roles":
            blocks.extend(_prd_contributor_blocks(document))
        elif document.document_type is DocumentType.PRD and definition.key == "key_dates":
            blocks.extend(_prd_milestone_blocks(document))
        elif definition.key not in hidden[document.document_type]:
            blocks.append(ExportBlock("text", definition.label, _clean(document.sections.get(definition.key))))
        if index == len(definitions) - 1:
            if document.document_type is DocumentType.BRD:
                blocks.extend(_brd_risk_blocks(document))
    return tuple(blocks)


def _validate_export_input(product: Product, document: ProductDocument) -> None:
    if document.document_type not in (DocumentType.BRD, DocumentType.PRD):
        raise DocumentExportError("Only saved BRDs and PRDs can be exported.")
    if product.id is None or document.id is None or document.product_id != product.id:
        raise DocumentExportError("Only a saved document associated with this product can be exported.")


def create_document_export(
    product: Product,
    document: ProductDocument,
    export_format: ExportFormat | str,
    *,
    generated_at: datetime | None = None,
) -> ExportResult:
    """Build one export fully in memory without database or provider access."""

    _validate_export_input(product, document)
    try:
        selected_format = ExportFormat(export_format)
    except (TypeError, ValueError) as error:
        raise DocumentExportError("Choose Word or PDF export.") from error
    timestamp = _utc_time(generated_at)
    try:
        if selected_format is ExportFormat.DOCX:
            content = _create_docx(product, document, timestamp)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        else:
            content = _create_pdf(product, document, timestamp)
            mime_type = "application/pdf"
    except DocumentExportError:
        raise
    except Exception as error:
        raise DocumentExportError("The document could not be exported safely. Please try again.") from error
    return ExportResult(
        content=content,
        filename=export_filename(product, document, selected_format),
        mime_type=mime_type,
        generated_at=timestamp,
    )


def _set_cell_margins(cell: object) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", 80), ("start", 120), ("bottom", 80), ("end", 120)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_run(run: object, *, size: float = 11, color: str = "000000", bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def _set_table_geometry(table: object, widths: tuple[float, ...]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_margins(cell)


def _add_docx_table(document: Document, block: ExportBlock) -> None:
    heading = document.add_paragraph(style="Heading 2")
    heading.paragraph_format.keep_with_next = True
    _set_run(heading.add_run(block.title), size=13, color=BLUE, bold=True)
    rows = tuple(block.content)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    _set_table_geometry(table, block.widths)
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            lines = str(value).splitlines() or [""]
            for line_index, line in enumerate(lines):
                if line_index:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                _set_run(
                    paragraph.add_run(line),
                    size=9.5,
                    color=INK if row_index == 0 else "000000",
                    bold=row_index == 0 or (column_index == 0 and len(rows[0]) == 2),
                )
            if row_index == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), LIGHT_FILL)
                cell._tc.get_or_add_tcPr().append(shading)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_page_field(paragraph: object) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    _set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, result, end):
        run._r.append(node)


def _normalize_docx_archive(raw: bytes) -> bytes:
    source = BytesIO(raw)
    output = BytesIO()
    with ZipFile(source) as input_zip, ZipFile(output, "w", ZIP_DEFLATED, compresslevel=9) as output_zip:
        for name in sorted(input_zip.namelist()):
            info = ZipInfo(name, date_time=(2000, 1, 1, 0, 0, 0))
            info.compress_type = ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            output_zip.writestr(info, input_zip.read(name))
    return output.getvalue()


def _create_docx(product: Product, document: ProductDocument, generated_at: datetime) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = section.right_margin = section.bottom_margin = section.left_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    _set_run(header.add_run(f"Product Manager Central · {document.document_type.value}"), size=9, color=MUTED, bold=True)
    footer = section.footer.paragraphs[0]
    _add_page_field(footer)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    _set_run(title.add_run(document.title), size=23, color="000000", bold=True)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    _set_run(subtitle.add_run(f"{document.document_type.value} · {_status_label(document)}"), size=14, color=MUTED)
    metadata = (
        ("Product", f"{product.name} (ID {product.id})"),
        ("Document", f"ID {document.id} · Version {document.version}"),
        ("Status", _status_label(document)),
        ("Generated", _display_time(generated_at)),
    )
    metadata_table = doc.add_table(rows=len(metadata), cols=2)
    metadata_table.style = "Table Grid"
    _set_table_geometry(metadata_table, (1.35, 5.15))
    for index, (label, value) in enumerate(metadata):
        for column, text in enumerate((label, value)):
            cell = metadata_table.cell(index, column)
            cell.text = ""
            _set_run(cell.paragraphs[0].add_run(text), size=10, bold=column == 0, color=INK if column == 0 else "000000")
            if column == 0:
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), LIGHT_FILL)
                cell._tc.get_or_add_tcPr().append(shading)
    doc.add_paragraph()

    for block in export_blocks(document):
        if block.kind == "group":
            paragraph = doc.add_paragraph(style="Heading 1")
            _set_run(paragraph.add_run(block.title), size=16, color=BLUE, bold=True)
        elif block.kind == "heading":
            paragraph = doc.add_paragraph(style="Heading 1")
            _set_run(paragraph.add_run(block.title), size=16, color=BLUE, bold=True)
        elif block.kind == "text":
            heading = doc.add_paragraph(style="Heading 2")
            _set_run(heading.add_run(block.title), size=13, color=BLUE, bold=True)
            text = str(block.content)
            paragraph = doc.add_paragraph()
            for index, line in enumerate(text.splitlines() or [""]):
                if index:
                    paragraph.add_run().add_break(WD_BREAK.LINE)
                _set_run(paragraph.add_run(line), size=11)
        elif block.kind == "table":
            _add_docx_table(doc, block)

    properties = doc.core_properties
    properties.title = document.title
    properties.subject = f"{document.document_type.value} export for {product.name}"
    properties.author = "Product Manager Central"
    properties.last_modified_by = "Product Manager Central"
    properties.keywords = ""
    properties.comments = ""
    naive_utc = generated_at.replace(tzinfo=None)
    properties.created = naive_utc
    properties.modified = naive_utc
    buffer = BytesIO()
    doc.save(buffer)
    return _normalize_docx_archive(buffer.getvalue())


def _register_pdf_font() -> str:
    candidates = (
        Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
        Path("/System/Library/Fonts/Supplemental/Arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    )
    for candidate in candidates:
        if candidate.is_file():
            name = "PMCUnicode"
            if name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(name, str(candidate)))
            return name
    raise DocumentExportError("PDF export needs a supported local Unicode font.")


class _PDFDocument(BaseDocTemplate):
    def __init__(self, buffer: BytesIO, *, title: str, status: str, generated_at: datetime) -> None:
        super().__init__(
            buffer,
            pagesize=LETTER,
            leftMargin=inch,
            rightMargin=inch,
            topMargin=inch,
            bottomMargin=inch,
            title=title,
            author="Product Manager Central",
            subject="Saved product document export",
            invariant=1,
        )
        self.export_title = title
        self.export_status = status
        self.generated_at = generated_at
        frame = Frame(self.leftMargin, self.bottomMargin, self.width, self.height, id="body")
        self.addPageTemplates(PageTemplate(id="pmc", frames=(frame,), onPage=self._decorate_page))

    def _decorate_page(self, canvas: object, doc: object) -> None:
        canvas.saveState()
        canvas.setFont("PMCUnicode", 8)
        canvas.setFillColor(colors.HexColor(f"#{MUTED}"))
        canvas.drawString(inch, 10.42 * inch, "Product Manager Central")
        canvas.drawRightString(7.5 * inch, 0.52 * inch, f"{self.export_status} · Page {doc.page}")
        canvas.restoreState()


def _pdf_text(value: object) -> str:
    return escape(str(value), quote=False).replace("\n", "<br/>")


def _create_pdf(product: Product, document: ProductDocument, generated_at: datetime) -> bytes:
    font_name = _register_pdf_font()
    styles = getSampleStyleSheet()
    body = ParagraphStyle(
        "PMCBody", parent=styles["BodyText"], fontName=font_name, fontSize=10,
        leading=12, spaceAfter=6, textColor=colors.black, alignment=TA_LEFT,
    )
    title = ParagraphStyle(
        "PMCTitle", parent=body, fontSize=22, leading=26, spaceAfter=4,
        textColor=colors.black,
    )
    subtitle = ParagraphStyle("PMCSubtitle", parent=body, fontSize=13, leading=16, spaceAfter=14, textColor=colors.HexColor(f"#{MUTED}"))
    h1 = ParagraphStyle("PMCH1", parent=body, fontSize=15, leading=18, spaceBefore=14, spaceAfter=7, textColor=colors.HexColor(f"#{BLUE}"), keepWithNext=True)
    h2 = ParagraphStyle("PMCH2", parent=body, fontSize=12, leading=15, spaceBefore=10, spaceAfter=5, textColor=colors.HexColor(f"#{BLUE}"), keepWithNext=True)
    cell = ParagraphStyle("PMCCell", parent=body, fontSize=8.5, leading=10.5, spaceAfter=0)
    cell_header = ParagraphStyle("PMCCellHeader", parent=cell, textColor=colors.HexColor(f"#{INK}"))

    buffer = BytesIO()
    pdf = _PDFDocument(buffer, title=document.title, status=_status_label(document), generated_at=generated_at)
    story: list[object] = [
        Paragraph(_pdf_text(document.title), title),
        Paragraph(_pdf_text(f"{document.document_type.value} · {_status_label(document)}"), subtitle),
    ]
    metadata = (
        ("Product", f"{product.name} (ID {product.id})"),
        ("Document", f"ID {document.id} · Version {document.version}"),
        ("Status", _status_label(document)),
        ("Generated", _display_time(generated_at)),
    )
    metadata_data = [[Paragraph(_pdf_text(label), cell_header), Paragraph(_pdf_text(value), cell)] for label, value in metadata]
    metadata_table = Table(metadata_data, colWidths=(1.35 * inch, 5.15 * inch), repeatRows=0, hAlign="LEFT")
    metadata_table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CCD3DB")),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor(f"#{LIGHT_FILL}")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6), ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    story.extend((metadata_table, Spacer(1, 8)))

    for block in export_blocks(document):
        if block.kind in {"group", "heading"}:
            story.append(Paragraph(_pdf_text(block.title), h1))
        elif block.kind == "text":
            story.append(Paragraph(_pdf_text(block.title), h2))
            story.append(Paragraph(_pdf_text(block.content), body))
        elif block.kind == "table":
            story.append(Paragraph(_pdf_text(block.title), h2))
            rows = tuple(block.content)
            table_data = [
                [Paragraph(_pdf_text(value), cell_header if row_index == 0 else cell) for value in row]
                for row_index, row in enumerate(rows)
            ]
            table = Table(
                table_data,
                colWidths=tuple(width * inch for width in block.widths),
                repeatRows=1,
                hAlign="LEFT",
                splitByRow=1,
            )
            table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#CCD3DB")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(f"#{LIGHT_FILL}")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.extend((table, Spacer(1, 5)))
    pdf.build(story)
    return buffer.getvalue()
