"""Focused Checkpoint 12 Word/PDF export tests."""

from datetime import datetime, timezone
from hashlib import sha256
from io import BytesIO
import os
from pathlib import Path
import re
import sqlite3
import tempfile
import unittest
from unittest.mock import patch
from zipfile import ZipFile

from docx import Document
from streamlit.testing.v1 import AppTest

from src.database import create_document, create_product, get_document, initialize_database
from src.document_export import (
    DocumentExportError,
    ExportFormat,
    create_document_export,
    export_blocks,
    export_filename,
)
from src.document_templates import document_template
from src.models import DocumentStatus, DocumentType
from tests.success_matrix_fixtures import complete_prd_agile_hierarchy, complete_success_matrix


APP_FILE = Path(__file__).resolve().parents[1] / "app.py"
FIXED_TIME = datetime(2026, 8, 14, 19, 0, tzinfo=timezone.utc)


def product_data(name: str = "Atlas Café / ../ Roadmap") -> dict[str, object]:
    return {
        "name": name,
        "description": "A planning workspace.",
        "target_users": "Product teams",
        "business_goal": "Improve measurable decisions.",
        "status": "planning",
        "customer_problem": "Requirements are fragmented.",
    }


def _sections(document_type: DocumentType, *, approved: bool = True) -> dict[str, str]:
    return {
        definition.key: (
            f"Saved {definition.label} with Unicode café 東京 and multiline:\n"
            "• First preserved statement\n• Second preserved statement"
            if approved else ""
        )
        for definition in document_template(document_type)
    }


def brd_hierarchy() -> list[dict[str, object]]:
    row: dict[str, object] = {"row_id": "brd-row-1", "position": 1}
    parent = None
    for level in ("epic", "capability", "feature", "user_story"):
        item_id = f"brd-{level}-1"
        row[f"{level}_id"] = item_id
        if parent is not None:
            row[f"{level}_parent_id"] = parent
        row[level] = f"Saved {level.replace('_', ' ')}"
        row[f"{level}_acceptance_criteria"] = [{
            "criterion_id": f"{item_id}-criterion-1",
            "position": 1,
            "text": f"The {level.replace('_', ' ')} is measurably complete.",
        }]
        parent = item_id
    return [row]


def document_data(product_id: int, document_type: DocumentType, *, approved: bool = True) -> dict[str, object]:
    data: dict[str, object] = {
        "product_id": product_id,
        "document_type": document_type,
        "title": f"Atlas {document_type.value} <script>alert(1)</script>",
        "version": "1.0/../../unsafe",
        "document_status": DocumentStatus.APPROVED if approved else DocumentStatus.DRAFT,
        "sections": _sections(document_type, approved=approved),
        "success_matrix": [],
        "agile_hierarchy": [],
        "contributors": [],
        "key_dates_milestones": [],
        "brd_hierarchy": [],
        "brd_risks": [],
    }
    if document_type is DocumentType.PRD and approved:
        data.update({
            "success_matrix": complete_success_matrix(),
            "agile_hierarchy": complete_prd_agile_hierarchy("export"),
            "contributors": [
                {"entry_id": "contributor-1", "contributor_name": "Zoë 李", "contributor_role": "Product Manager"},
                {"entry_id": "contributor-2", "contributor_name": "Renée", "contributor_role": "Analytics Lead"},
            ],
            "key_dates_milestones": [
                {"entry_id": "milestone-1", "date": "2026-09-01", "milestone": "Beta begins"},
                {"entry_id": "milestone-2", "date": "2026-10-15", "milestone": "Launch review"},
            ],
        })
    elif document_type is DocumentType.BRD and approved:
        data.update({
            "brd_hierarchy": brd_hierarchy(),
            "brd_risks": [{
                "entry_id": "risk-1", "business_risk": "Low adoption",
                "mitigation_strategy": "Run a guided beta",
            }],
        })
    return data


class ExportTestCase(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.database = Path(self.temp.name) / "checkpoint12.db"
        initialize_database(self.database)
        self.product = create_product(product_data(), self.database)

    def create(self, document_type: DocumentType, *, approved: bool = True):
        return create_document(document_data(self.product.id, document_type, approved=approved), self.database)


class FilenameAndBoundaryTests(ExportTestCase):
    def test_filename_is_deterministic_descriptive_and_path_safe(self):
        document = self.create(DocumentType.PRD)
        filename = export_filename(self.product, document, ExportFormat.DOCX)
        self.assertEqual(filename, export_filename(self.product, document, "docx"))
        self.assertRegex(filename, rf"^pmc-atlas-cafe-roadmap-prd-{document.id}-v1-0-unsafe\.docx$")
        self.assertNotIn("..", filename)
        self.assertNotIn("/", filename)
        self.assertNotIn("\\", filename)
        self.assertFalse(filename.startswith("."))

    def test_invalid_format_and_unsaved_or_cross_product_documents_fail_safely(self):
        document = self.create(DocumentType.PRD)
        with self.assertRaisesRegex(DocumentExportError, "Choose Word or PDF"):
            create_document_export(self.product, document, "html")
        other = create_product(product_data("Other"), self.database)
        with self.assertRaisesRegex(DocumentExportError, "associated"):
            create_document_export(other, document, "pdf")

    def test_export_requires_no_openai_key_or_provider_call(self):
        document = self.create(DocumentType.BRD)
        with patch.dict(os.environ, {"OPENAI_API_KEY": ""}), patch(
            "src.ai_service.OpenAIService.from_environment"
        ) as provider:
            result = create_document_export(self.product, document, "pdf", generated_at=FIXED_TIME)
        self.assertTrue(result.content.startswith(b"%PDF"))
        provider.assert_not_called()

    def test_export_does_not_mutate_database_or_saved_record(self):
        document = self.create(DocumentType.PRD)
        before_bytes = self.database.read_bytes()
        before_record = get_document(document.id, self.database)
        create_document_export(self.product, document, "docx", generated_at=FIXED_TIME)
        create_document_export(self.product, document, "pdf", generated_at=FIXED_TIME)
        self.assertEqual(self.database.read_bytes(), before_bytes)
        self.assertEqual(get_document(document.id, self.database), before_record)


class OrderedContentTests(ExportTestCase):
    def test_prd_blocks_preserve_template_order_and_all_structured_collections(self):
        document = self.create(DocumentType.PRD)
        blocks = export_blocks(document)
        titles = [block.title for block in blocks]
        groups = [block.title for block in blocks if block.kind == "group"]
        self.assertEqual(groups, list(dict.fromkeys(item.group for item in document_template(DocumentType.PRD))))
        for title in (
            "Contributors and Roles", "Key Dates and Milestones",
            "Structured Agile Hierarchy", "PRD Success Matrix",
        ):
            self.assertIn(title, titles)
        self.assertLess(titles.index("Structured Agile Hierarchy"), titles.index("5. Nonfunctional Requirements"))
        self.assertLess(titles.index("PRD Success Matrix"), titles.index("8. Assumptions, Constraints, and Dependencies"))

    def test_brd_blocks_include_readable_hierarchy_criteria_risks_and_mitigation(self):
        document = self.create(DocumentType.BRD)
        combined = "\n".join(
            str(value)
            for block in export_blocks(document)
            for value in (block.title, block.content)
        )
        for expected in (
            "BRD Agile Hierarchy", "Saved epic", "Epic acceptance criteria",
            "Saved capability", "Saved feature", "Saved user story",
            "Business Risk and Mitigation Strategy", "Low adoption", "Run a guided beta",
        ):
            self.assertIn(expected, combined)

    def test_legacy_derived_rows_do_not_duplicate_saved_section_text(self):
        data = document_data(self.product.id, DocumentType.BRD, approved=False)
        data["sections"].update({
            "epics": "Legacy epic", "capabilities": "Legacy capability",
            "features": "Legacy feature", "user_stories": "Legacy story",
            "acceptance_criteria": "Legacy criterion", "business_risks": "Legacy risk",
            "mitigation_strategies": "Legacy mitigation",
        })
        for key in ("brd_hierarchy", "brd_risks"):
            data.pop(key)
        document = create_document(data, self.database)
        combined = "\n".join(str(block.content) for block in export_blocks(document))
        for text in ("Legacy epic", "Legacy capability", "Legacy feature", "Legacy story", "Legacy risk", "Legacy mitigation"):
            self.assertEqual(combined.count(text), 1)

    def test_distinct_legacy_content_is_labeled_and_preserved(self):
        document = self.create(DocumentType.PRD)
        titles = [block.title for block in export_blocks(document)]
        self.assertIn("Contributors and roles — preserved legacy content", titles)
        self.assertIn("Key dates — preserved legacy content", titles)
        self.assertIn("Milestones — preserved legacy content", titles)

    def test_blank_draft_sections_and_empty_structures_are_labeled(self):
        document = self.create(DocumentType.PRD, approved=False)
        combined = "\n".join(str(block.content) for block in export_blocks(document))
        self.assertIn("Not provided", combined)
        self.assertIn("No structured hierarchy entries provided", combined)
        self.assertIn("No Success Matrix entries provided", combined)


class WordExportTests(ExportTestCase):
    def test_docx_opens_and_preserves_unicode_multiline_metadata_and_status(self):
        document = self.create(DocumentType.PRD)
        result = create_document_export(self.product, document, "docx", generated_at=FIXED_TIME)
        opened = Document(BytesIO(result.content))
        text = "\n".join(paragraph.text for paragraph in opened.paragraphs)
        table_text = "\n".join(cell.text for table in opened.tables for row in table.rows for cell in row.cells)
        for expected in ("Atlas PRD <script>alert(1)</script>", "Approved", "2026-08-14 19:00 UTC", "café 東京", "Zoë 李", "Outcome completion rate"):
            self.assertIn(expected, text + table_text)
        self.assertIn("First preserved statement\n• Second preserved statement", text + table_text)

    def test_docx_is_macro_free_has_no_external_relationships_and_scrubs_personal_metadata(self):
        document = self.create(DocumentType.BRD)
        result = create_document_export(self.product, document, "docx", generated_at=FIXED_TIME)
        with ZipFile(BytesIO(result.content)) as archive:
            names = archive.namelist()
            self.assertFalse(any(name.endswith("vbaProject.bin") for name in names))
            self.assertNotIn("docProps/custom.xml", names)
            relationships = "\n".join(
                archive.read(name).decode("utf-8")
                for name in names if name.endswith(".rels")
            )
            self.assertNotIn('TargetMode="External"', relationships)
            core = archive.read("docProps/core.xml").decode("utf-8")
            self.assertIn("Product Manager Central", core)
            self.assertNotIn(str(Path.home()), core)

    def test_docx_encodes_letter_geometry_fixed_tables_and_page_field(self):
        document = self.create(DocumentType.PRD)
        result = create_document_export(self.product, document, "docx", generated_at=FIXED_TIME)
        with ZipFile(BytesIO(result.content)) as archive:
            xml = archive.read("word/document.xml").decode("utf-8")
            footer = "\n".join(
                archive.read(name).decode("utf-8") for name in archive.namelist() if name.startswith("word/footer")
            )
        self.assertRegex(xml, r'<w:pgSz[^>]+w:w="12240"[^>]+w:h="15840"')
        self.assertIn('w:tblLayout w:type="fixed"', xml)
        self.assertIn('w:tblW w:type="dxa" w:w="9360"', xml)
        self.assertIn(" PAGE ", footer)

    def test_docx_is_byte_deterministic_for_same_saved_input_and_time(self):
        document = self.create(DocumentType.PRD)
        first = create_document_export(self.product, document, "docx", generated_at=FIXED_TIME)
        second = create_document_export(self.product, document, "docx", generated_at=FIXED_TIME)
        self.assertEqual(first.content, second.content)
        self.assertEqual(sha256(first.content).hexdigest(), sha256(second.content).hexdigest())


class PDFExportTests(ExportTestCase):
    def test_pdf_signature_trailer_pages_and_in_memory_mime(self):
        document = self.create(DocumentType.PRD)
        result = create_document_export(self.product, document, "pdf", generated_at=FIXED_TIME)
        self.assertTrue(result.content.startswith(b"%PDF-"))
        self.assertIn(b"%%EOF", result.content[-1024:])
        self.assertGreater(len(re.findall(rb"/Type\s*/Page\b", result.content)), 1)
        self.assertEqual(result.mime_type, "application/pdf")
        self.assertTrue(result.filename.endswith(".pdf"))

    def test_pdf_is_byte_deterministic_for_same_saved_input_and_time(self):
        document = self.create(DocumentType.BRD)
        first = create_document_export(self.product, document, "pdf", generated_at=FIXED_TIME)
        second = create_document_export(self.product, document, "pdf", generated_at=FIXED_TIME)
        self.assertEqual(first.content, second.content)

    def test_pdf_long_content_flows_to_multiple_pages_without_size_limit_failure(self):
        data = document_data(self.product.id, DocumentType.PRD, approved=True)
        data["sections"]["product_overview"] = "\n".join(
            f"Line {index}: a deliberately long saved product statement." for index in range(1, 120)
        )
        document = create_document(data, self.database)
        result = create_document_export(self.product, document, "pdf", generated_at=FIXED_TIME)
        self.assertGreater(len(re.findall(rb"/Type\s*/Page\b", result.content)), 5)


class StreamlitDownloadTests(ExportTestCase):
    def test_saved_preview_exposes_word_and_pdf_downloads_without_provider(self):
        document = self.create(DocumentType.PRD, approved=False)
        original_database = os.environ.get("PMC_DATABASE_FILE")
        os.environ["PMC_DATABASE_FILE"] = str(self.database)
        self.addCleanup(
            lambda: os.environ.pop("PMC_DATABASE_FILE", None)
            if original_database is None
            else os.environ.__setitem__("PMC_DATABASE_FILE", original_database)
        )
        with patch("src.ai_service.OpenAIService.from_environment") as provider:
            app = AppTest.from_file(APP_FILE, default_timeout=8).run()
            app.radio[0].set_value("View Products").run()
            app.button(key="view_product_selector_preview_document").click().run()
        self.assertEqual(list(app.exception), [])
        rendered = "\n".join(str(element.value) for element in (*app.markdown, *app.caption))
        self.assertIn("Download saved document", rendered)
        self.assertIn("Downloading does not change the document or require an API key", rendered)
        provider.assert_not_called()


if __name__ == "__main__":
    unittest.main()
